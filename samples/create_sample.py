from docx import Document

doc = Document()
doc.add_heading('Legacy System Report', 0)

doc.add_paragraph('The current system is built on a monolithic architecture that is difficult to scale. We have observed significant latency during peak hours.')

doc.add_heading('Current Issues', level=1)
doc.add_paragraph('1. Slow database queries due to lack of indexing.')
doc.add_paragraph('2. Security vulnerabilities in outdated dependencies.')
doc.add_paragraph('3. High maintenance costs.')

doc.add_heading('Proposed Solution', level=1)
doc.add_paragraph('We propose a complete rewrite using microservices. This will allow us to scale components independently.')

doc.save('sample.docx')
print("Created sample.docx")
