import uuid
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
import re

class DocxHandler:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = Document(file_path)
        self.id_map = {} # Maps UUID -> (Type, Object)
        self.structure = []

    def extract_structure(self):
        """
        Iterates through the document and creates a JSON-serializable structure 
        (HTML preview + metadata) while populating the ID map.
        """
        self.id_map = {}
        self.structure = []
        html_parts = []

        # Helper to process block
        def process_block(block, block_type):
            block_id = str(uuid.uuid4())
            self.id_map[block_id] = {"type": block_type, "obj": block}
            
            if block_type == "paragraph":
                text = block.text
                style = block.style.name
                # Simple style mapping to HTML classes
                css_class = "mb-2"
                if "Heading 1" in style: css_class = "text-3xl font-bold mt-6 mb-4"
                elif "Heading 2" in style: css_class = "text-2xl font-bold mt-5 mb-3"
                elif "Heading 3" in style: css_class = "text-xl font-bold mt-4 mb-2"
                elif "List" in style: css_class = "list-disc ml-6 mb-1"
                elif "Normal" in style: css_class = "mb-2 leading-relaxed"
                
                # Basic run formatting (bold/italic) for preview
                # Note: This is a simplified preview. For perfect fidelity we rely on the docx object.
                inner_html = ""
                for run in block.runs:
                    run_text = run.text
                    if run.bold: run_text = f"<b>{run_text}</b>"
                    if run.italic: run_text = f"<i>{run_text}</i>"
                    inner_html += run_text
                
                if not inner_html.strip(): inner_html = "&nbsp;" # Preserve empty lines

                html = f'<div id="{block_id}" class="docx-section p-2 hover:bg-blue-50 cursor-pointer transition-colors rounded {css_class}" data-type="paragraph">{inner_html}</div>'
                
                self.structure.append({
                    "id": block_id,
                    "type": "paragraph",
                    "text": text,
                    "style": style
                })
                return html

            elif block_type == "table":
                # Simplified table preview
                rows_html = ""
                for row in block.rows:
                    cells_html = ""
                    for cell in row.cells:
                        # We treat cells as containers of paragraphs usually, but for simple preview:
                        cell_text = cell.text
                        cells_html += f'<td class="border p-2">{cell_text}</td>'
                    rows_html += f"<tr>{cells_html}</tr>"
                
                html = f'<div id="{block_id}" class="docx-section my-4 overflow-x-auto" data-type="table"><table class="w-full border-collapse border">{rows_html}</table></div>'
                self.structure.append({
                    "id": block_id,
                    "type": "table",
                    "text": "[Table]"
                })
                return html
            return ""

        # Iterate through document body
        # python-docx doesn't have a unified iterator for paragraphs and tables in order easily
        # We will use iter_block_items approach if available, or just paragraphs for MVP if mixed
        # For MVP, let's stick to paragraphs as they are 90% of content. 
        # To do it right, we access element.body.iter_children() (requires lxml access)
        
        # Robust iteration:
        for element in self.doc.element.body:
            if element.tag.endswith('p'):
                # It's a paragraph
                # Find the paragraph object wrapping this element
                # This is O(N) search if not careful, but python-docx creates wrappers on fly.
                # Better: iterate doc.paragraphs and doc.tables and sort by element location? 
                # No, that's slow.
                # Let's just iterate doc.paragraphs for now. 
                # If the user needs tables mixed in perfectly, we need a more advanced parser.
                pass

        # SIMPLIFIED APPROACH FOR MVP: Just Paragraphs. 
        # Tables often break flow in simple editors. Let's support Paragraphs fully first.
        for p in self.doc.paragraphs:
            if p.text.strip() or True: # Include empty lines for spacing
                html_parts.append(process_block(p, "paragraph"))

        return "".join(html_parts)

    def update_section(self, section_id: str, new_text: str):
        if section_id not in self.id_map:
            return False
        
        item = self.id_map[section_id]
        if item["type"] == "paragraph":
            p = item["obj"]
            p.clear() # Remove old content
            
            # Simple Markdown Parser for Bold/Italic
            # We split by ** for bold, then * for italic
            # This is a basic parser.
            
            # Regex for **bold**
            parts = re.split(r'(\*\*.*?\*\*)', new_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    # Split for *italic*
                    subparts = re.split(r'(\*.*?\*)', part)
                    for subpart in subparts:
                        if subpart.startswith('*') and subpart.endswith('*'):
                            run = p.add_run(subpart[1:-1])
                            run.italic = True
                        else:
                            p.add_run(subpart)
            return True
        return False

    def save(self, output_path: str):
        self.doc.save(output_path)
